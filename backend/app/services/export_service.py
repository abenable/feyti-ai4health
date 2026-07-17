"""Markdown to DOCX export."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Inches


def _set_margins(doc: Document) -> None:
    sections = doc.sections
    if not sections:
        return
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_inline(paragraph, text: str) -> None:
    """Render a subset of inline Markdown (**bold**, [text](url)) into runs."""
    # Turn links into their display text first; a DOCX paragraph can't be a
    # hyperlink without extra plumbing, and the text is what a reviewer reads.
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    # Split on **bold** spans, keeping the delimiters' content.
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1  # odd chunks are the captured bold spans


def markdown_to_docx(markdown: str, title: str) -> bytes:
    """Convert simple Markdown to a DOCX file; return the bytes."""
    doc = Document()
    _set_margins(doc)

    # Cover / header title.
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = doc.styles["Title"].font.size
        p.alignment = 1  # center

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        # ATX headings.
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            style = f"Heading {min(level, 9)}"  # docx ships Heading 1..9
            doc.add_paragraph(heading_match.group(2).strip(), style=style)
            continue

        # Ordered list items (1. 2. …).
        ordered_match = re.match(r"^\d+[.)]\s+(.*)$", line)
        if ordered_match:
            _add_inline(doc.add_paragraph(style="List Number"), ordered_match.group(1).strip())
            continue

        # Bullet lines (start with -, *, +).
        bullet_match = re.match(r"^[-*+]\s+(.*)$", line)
        if bullet_match:
            _add_inline(doc.add_paragraph(style="List Bullet"), bullet_match.group(1).strip())
            continue

        _add_inline(doc.add_paragraph(), line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if __name__ == "__main__":  # ponytail self-check
    sample = (
        "# Stability Data\n\n"
        "This section contains stability data.\n\n"
        "## Conditions\n\n"
        "- 25C / 60% RH\n"
        "- 40C / 75% RH\n\n"
        "Conclusion: acceptable."
    )
    docx_bytes = markdown_to_docx(sample, "Stability Data")
    assert docx_bytes, "markdown_to_docx returned empty bytes"
    assert b"PK" == docx_bytes[:2], "output does not look like a ZIP/DOCX"
    print("OK — markdown_to_docx produces non-empty DOCX bytes")
